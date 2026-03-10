from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import TYPE_CHECKING, Any, cast

from rag_challenge.config.settings import IngestionSettings
from rag_challenge.models import DocType, DocumentSection, ParsedDocument, ProvidedChunk

if TYPE_CHECKING:
    from docling.document_converter import DocumentConverter

logger = logging.getLogger(__name__)

_SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".text", ".json"}

_STATUTE_PATTERNS = re.compile(
    r"(section\s+\d|article\s+\d|§\s*\d|chapter\s+\d|title\s+\d|act\b|statute|regulation|code)",
    re.IGNORECASE,
)
_CASE_PATTERNS = re.compile(
    r"( v\. | vs\. |plaintiff|defendant|court|judgment|opinion|ruling|appellant|respondent)",
    re.IGNORECASE,
)
_CONTRACT_PATTERNS = re.compile(
    r"(agreement|contract|party|parties|hereby|whereas|witnesseth|covenant|indemnif|termination)",
    re.IGNORECASE,
)
_JURISDICTION_PATTERNS = re.compile(
    r"\b(U\.?S\.?A?|United States|United Kingdom|U\.?K\.?|European Union|E\.?U\.?|Australia|Singapore|Ireland)\b",
    re.IGNORECASE,
)

_HEADING_RE = re.compile(
    r"^(#{1,6}\s+.+"
    r"|(?:ARTICLE|Article|SECTION|Section|CHAPTER|Chapter|PART|Part)\s+\w+.*"
    r"|\d+(?:\.\d+)*\s+[A-Z].+"
    r"|[A-Z][A-Z\s]{4,})$"
)


class DocumentParser:
    """Parse legal files into structured `ParsedDocument` objects."""

    def __init__(
        self,
        *,
        pdf_text_min_chars: int | None = None,
        pdf_text_min_words: int | None = None,
    ) -> None:
        ingestion_settings = IngestionSettings()
        self._pdf_text_min_chars = int(
            pdf_text_min_chars
            if pdf_text_min_chars is not None
            else ingestion_settings.parser_pdf_text_min_chars
        )
        self._pdf_text_min_words = int(
            pdf_text_min_words
            if pdf_text_min_words is not None
            else ingestion_settings.parser_pdf_text_min_words
        )
        self._docling_converter: DocumentConverter | None = None

    def list_supported_files(self, dir_path: str | Path) -> list[Path]:
        root = Path(dir_path)
        if not root.is_dir():
            raise NotADirectoryError(f"Not a directory: {root}")

        return sorted(p for p in root.rglob("*") if p.is_file() and p.suffix.lower() in _SUPPORTED_EXTENSIONS)

    def parse_file(self, path: str | Path) -> ParsedDocument:
        file_path = Path(path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        if not file_path.is_file():
            raise ValueError(f"Path is not a file: {file_path}")

        suffix = file_path.suffix.lower()
        if suffix not in _SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {suffix}")

        # IMPORTANT: doc IDs must be stable across environments/paths (evaluation runs in a different filesystem).
        # Use the filename stem so `pdf_id_page` identifiers can be reproduced deterministically.
        doc_id = self._generate_doc_id(file_path)
        if suffix == ".json":
            return self._parse_json_document(file_path, fallback_doc_id=doc_id)

        raw_text = ""
        pdf_pages: list[str] | None = None
        if suffix == ".pdf":
            raw_text, pdf_pages = self._read_pdf(file_path)
        else:
            raw_text = self._read_file(file_path)
        if not raw_text.strip():
            logger.warning("Empty or unparseable file: %s", file_path)
            return ParsedDocument(
                doc_id=doc_id,
                title=self._fallback_title_from_path(file_path),
                doc_type=DocType.OTHER,
                source_path=str(file_path),
                full_text="",
                sections=[],
            )

        title = self._extract_title(raw_text, file_path)
        doc_type = self._classify_doc_type(raw_text)
        jurisdiction = self._extract_jurisdiction(raw_text)
        if pdf_pages is not None and pdf_pages:
            # Competition grounding uses page IDs like `pdf_id_page`.
            sections = [
                DocumentSection(
                    heading=f"Page {idx + 1}",
                    section_path=f"page:{idx + 1}",
                    text=page_text,
                    level=0,
                )
                for idx, page_text in enumerate(pdf_pages)
                if page_text.strip()
            ]
        else:
            sections = self._extract_sections(raw_text)

        return ParsedDocument(
            doc_id=doc_id,
            title=title,
            doc_type=doc_type,
            jurisdiction=jurisdiction,
            source_path=str(file_path),
            full_text=raw_text,
            sections=sections,
        )

    def parse_directory(self, dir_path: str | Path) -> list[ParsedDocument]:
        files = self.list_supported_files(dir_path)
        parsed: list[ParsedDocument] = []
        for file_path in files:
            try:
                parsed.append(self.parse_file(file_path))
            except Exception:
                logger.warning("Failed to parse %s; skipping", file_path, exc_info=True)
        return parsed

    def _read_file(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(path)
        if suffix == ".docx":
            return self._parse_docx(path)
        return path.read_text(encoding="utf-8", errors="replace")

    def _read_pdf(self, path: Path) -> tuple[str, list[str]]:
        pages = self._parse_pdf_pymupdf_pages(path)
        fast_text = "\n\n".join(text for text in pages if text).strip()
        if self._is_pdf_text_sufficient(fast_text):
            return fast_text, pages

        if fast_text.strip():
            logger.info(
                "PyMuPDF text extraction too short/low quality for %s (chars=%d words=%d); falling back to Docling",
                path,
                len(fast_text),
                len(fast_text.split()),
            )

        docling_text = self._parse_pdf_docling(path)
        # Docling may lose page boundaries; treat as a single-page section.
        merged = (docling_text or fast_text).strip()
        return merged, [merged] if merged else []

    def _parse_json_document(self, path: Path, *, fallback_doc_id: str) -> ParsedDocument:
        raw = path.read_text(encoding="utf-8", errors="replace")
        data_obj: object = json.loads(raw)
        if not isinstance(data_obj, dict):
            raise ValueError(f"JSON document must be an object: {path}")
        data = cast("dict[str, object]", data_obj)

        stable_doc_id_obj = data.get("doc_id")
        stable_doc_id = str(stable_doc_id_obj) if stable_doc_id_obj else fallback_doc_id
        title = str(data.get("title") or path.stem)
        jurisdiction = str(data.get("jurisdiction") or "")
        doc_type = self._coerce_doc_type(data.get("doc_type"))
        provided_chunks = self._parse_provided_chunks(data.get("chunks"))

        full_text = "\n\n".join(chunk.text for chunk in provided_chunks)
        sections = [
            DocumentSection(
                heading=chunk.section_path or "",
                section_path=chunk.section_path,
                text=chunk.text,
                level=1 if chunk.section_path else 0,
            )
            for chunk in provided_chunks
        ]

        return ParsedDocument(
            doc_id=stable_doc_id,
            title=title,
            doc_type=doc_type,
            jurisdiction=jurisdiction,
            source_path=str(path),
            full_text=full_text,
            sections=sections,
            provided_chunks=provided_chunks,
            metadata={"source_format": "json_prechunked"},
        )

    @staticmethod
    def _coerce_doc_type(value: object) -> DocType:
        if isinstance(value, str):
            raw = value.strip().lower()
            for doc_type in DocType:
                if raw == doc_type.value:
                    return doc_type
        return DocType.OTHER

    @staticmethod
    def _parse_provided_chunks(chunks_obj: object) -> list[ProvidedChunk]:
        if not isinstance(chunks_obj, list):
            return []
        chunks: list[ProvidedChunk] = []
        for item_obj in cast("list[object]", chunks_obj):
            if not isinstance(item_obj, dict):
                continue
            item = cast("dict[str, object]", item_obj)
            chunk_id_obj = item.get("chunk_id")
            text_obj = item.get("text")
            if not isinstance(chunk_id_obj, str) or not chunk_id_obj.strip():
                continue
            if not isinstance(text_obj, str) or not text_obj.strip():
                continue
            citations = item.get("citations")
            anchors = item.get("anchors")
            citation_values = (
                [str(v) for v in cast("list[object]", citations)] if isinstance(citations, list) else []
            )
            anchor_values = [str(v) for v in cast("list[object]", anchors)] if isinstance(anchors, list) else []
            chunks.append(
                ProvidedChunk(
                    chunk_id=chunk_id_obj.strip(),
                    text=text_obj,
                    section_path=str(item.get("section_path") or ""),
                    citations=citation_values,
                    anchors=anchor_values,
                )
            )
        return chunks

    def _parse_pdf(self, path: Path) -> str:
        text, _pages = self._read_pdf(path)
        return text

    def _parse_pdf_pymupdf(self, path: Path) -> str:
        pages = self._parse_pdf_pymupdf_pages(path)
        return "\n\n".join(text for text in pages if text).strip()

    def _parse_pdf_pymupdf_pages(self, path: Path) -> list[str]:
        try:
            import fitz  # pyright: ignore[reportMissingImports,reportMissingTypeStubs]

            pages: list[str] = []
            fitz_any = cast("Any", fitz)
            with fitz_any.open(str(path)) as pdf_obj:
                for page_obj in cast("list[object]", list(pdf_obj)):
                    page_any = cast("Any", page_obj)
                    page_text_obj = page_any.get_text("text")
                    if not isinstance(page_text_obj, str):
                        continue
                    page_text = page_text_obj
                    if page_text:
                        pages.append(page_text.strip())
            return pages
        except Exception:
            logger.info("PyMuPDF failed for %s; falling back to Docling", path, exc_info=True)
            return []

    def _parse_pdf_docling(self, path: Path) -> str:
        try:
            if self._docling_converter is None:
                from docling.datamodel.base_models import InputFormat
                from docling.datamodel.pipeline_options import PdfPipelineOptions
                from docling.document_converter import DocumentConverter, PdfFormatOption

                pipeline_options = PdfPipelineOptions(do_ocr=True)
                self._docling_converter = DocumentConverter(
                    format_options={
                        InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options),
                    }
                )

            result = self._docling_converter.convert(str(path))
            return str(result.document.export_to_markdown())
        except Exception:
            logger.warning("Docling failed for %s", path, exc_info=True)
            return ""

    def _is_pdf_text_sufficient(self, text: str) -> bool:
        stripped = text.strip()
        if not stripped:
            return False

        word_count = len(stripped.split())
        if len(stripped) < self._pdf_text_min_chars and word_count < self._pdf_text_min_words:
            return False

        visible_chars = sum(1 for char in stripped if char.isalnum())
        density = visible_chars / max(len(stripped), 1)
        return density >= 0.2

    def _parse_docx(self, path: Path) -> str:
        try:
            import docx

            document = docx.Document(str(path))
            paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception:
            logger.warning("DOCX parse failed for %s", path, exc_info=True)
            return ""

    @staticmethod
    def _generate_doc_id(path: Path) -> str:
        # Stable across environments; also matches `pdf_id` used in `pdf_id_page` identifiers.
        return path.stem

    @staticmethod
    def _fallback_title_from_path(path: Path) -> str:
        return path.stem.replace("_", " ").replace("-", " ").strip().title()

    @classmethod
    def _extract_title(cls, text: str, path: Path) -> str:
        stripped = text.strip()
        if not stripped:
            return cls._fallback_title_from_path(path)

        lines = [line.strip() for line in stripped.splitlines() if line.strip()]
        if not lines:
            return cls._fallback_title_from_path(path)

        first = lines[0]
        if first.startswith("#"):
            return first.lstrip("#").strip()
        if len(first) <= 200:
            return first
        return cls._fallback_title_from_path(path)

    @staticmethod
    def _classify_doc_type(text: str) -> DocType:
        sample = text[:5000]
        scores = {
            DocType.STATUTE: len(_STATUTE_PATTERNS.findall(sample)),
            DocType.CASE_LAW: len(_CASE_PATTERNS.findall(f" {sample} ")),
            DocType.CONTRACT: len(_CONTRACT_PATTERNS.findall(sample)),
        }
        best = max(scores, key=lambda doc_type: scores[doc_type])
        return best if scores[best] >= 2 else DocType.OTHER

    @staticmethod
    def _extract_jurisdiction(text: str) -> str:
        sample = text[:3000]
        match = _JURISDICTION_PATTERNS.search(sample)
        if match is None:
            return ""

        raw = match.group(0).strip().rstrip(".")
        normalized = {
            "U.S": "US",
            "U.S.A": "US",
            "USA": "US",
            "US": "US",
            "United States": "US",
            "U.K": "UK",
            "UK": "UK",
            "United Kingdom": "UK",
            "E.U": "EU",
            "EU": "EU",
            "European Union": "EU",
            "Australia": "AU",
            "Singapore": "SG",
            "Ireland": "IE",
        }
        return normalized.get(raw, raw)

    @staticmethod
    def _extract_sections(text: str) -> list[DocumentSection]:
        lines = text.splitlines()
        sections: list[DocumentSection] = []
        current_heading = ""
        current_level = 0
        current_lines: list[str] = []
        heading_path: list[str] = []

        def flush_current() -> None:
            section_text = "\n".join(current_lines).strip()
            if not section_text:
                return
            sections.append(
                DocumentSection(
                    heading=current_heading,
                    section_path=" > ".join(heading_path) if heading_path else current_heading,
                    text=section_text,
                    level=current_level,
                )
            )

        for raw_line in lines:
            line = raw_line.strip()
            if _HEADING_RE.match(line):
                if current_lines:
                    flush_current()

                current_heading = line.lstrip("#").strip()
                level = len(line) - len(line.lstrip("#")) if line.startswith("#") else 1
                current_level = level

                while heading_path and len(heading_path) >= level:
                    heading_path.pop()
                heading_path.append(current_heading)
                current_lines = []
                continue

            current_lines.append(raw_line)

        if current_lines:
            flush_current()

        if sections:
            return sections

        fallback_text = text.strip()
        if not fallback_text:
            return []
        return [
            DocumentSection(
                heading="",
                section_path="",
                text=fallback_text,
                level=0,
            )
        ]
